# _*_ coding:utf-8 _*_
#@Author = zhoupengfei
#@Time = 2018/9/20 9:03
from docx import Document
from docx.shared import Pt
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import demjson,requests
import time,os,json,re
from microserver.core.handlers import handlers
from microserver.utils import decorator
from microserver.conf import settings
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



class LawopinionHandler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(LawopinionHandler, self).prepare()
        self.law_comment = self.get_argument("law_comment")
        self.action_comment = self.get_argument("action_comment")
        self.law = self.get_argument("law")
        self.law_list = eval(self.law)
        self.cause = self.get_argument("cause")
        self.cause_list = eval(self.cause)
    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'znys')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir, '法律意见书_{}.docx'.format(now),)
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/znys/{}'.format('法律意见书_{}.docx'.format(now))
        return result

    def generate_lawapply(self,filepath):
        document = Document()
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置
        run_law_word = p.add_run(u"法律意见书")
        run_law_word.font.name = u'宋体'
        run_law_word.font.size = Pt(28)

        la_a = document.add_paragraph()
        run_lwa = la_a.add_run(u'    您好，我们根据您填写的信息出具了一份法律意见书，请查看。')
        run_lwa.font.name = u'宋体'
        run_lwa.font.size = Pt(13)
        ra = run_lwa._element
        ra.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        f1 = run_law_word._element
        f1.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        law_font= document.add_paragraph()
        run_law = law_font.add_run(u'法律意见')
        run_law.bold = True
        run_law.font.name = u'宋体'
        run_law.font.size = Pt(14)
        r = run_law._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        paragraph = document.add_paragraph()
        str_law = self.law_comment.replace('\n', '\n    ')
        run = paragraph.add_run(u'    '+str_law)
        run.font.name = u'宋体'
        run.font.size = Pt(13)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        law_font = document.add_paragraph()
        run_law = law_font.add_run(u'行动建议')
        run_law.bold = True
        run_law.font.name = u'宋体'
        run_law.font.size = Pt(14)
        r = run_law._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        action_font = document.add_paragraph()
        str_obj = self.action_comment.replace('\n', '\n    ')
        run_action = action_font.add_run(u'    '+str_obj)

        run_action.font.name = u'宋体'
        run_action.font.size = Pt(13)
        f = run_action._element
        f.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        la_font = document.add_paragraph()
        run_lw = la_font.add_run(u'法律法规')
        run_lw.bold = True
        run_lw.font.name = u'宋体'
        run_lw.font.size = Pt(14)
        r = run_lw._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        for law in self.law_list:
            self.title = law.setdefault('title', '')
            self.content = law.setdefault('content', '')
            title_font = document.add_paragraph()
            content_font = document.add_paragraph()
            run_title = title_font.add_run(self.title)
            run_content = content_font.add_run(self.content)
            run_title.font.name = u'宋体'
            run_title.font.size = Pt(13)
            f = run_title._element
            f.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            run_content.font.name = u'宋体'
            run_content.font.size = Pt(13)
            f = run_content._element
            f.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        l_font = document.add_paragraph()
        run_l = l_font.add_run(u'类案参考')
        run_l.bold = True
        run_l.font.name = u'宋体'
        run_l.font.size = Pt(14)
        r = run_l._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        for case in self.cause_list:
            case_font = document.add_paragraph()
            run_case = case_font.add_run('《'+ case +'》')
            run_case.font.name = u'宋体'
            run_case.font.size = Pt(13)
            r = run_case._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        para = document.add_paragraph()
        # 设置字号
        run_hou = para.add_run(u'重要提示： 结果仅供参考，具体案件以法院审理为准。')
        run_hou.font.size = Pt(14)
        # 保存文件
        document.save(filepath)


class RiskassessmentHandler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(RiskassessmentHandler, self).prepare()
        self.legalFare = self.get_argument("legalFare")
        self.mediateFare = self.get_argument("mediateFare")
        self.mediateTime = self.get_argument("mediateTime")
        self.trialTime = self.get_argument("trialTime")
        self.tiaojie = self.get_argument("tiaojie","")
        self.susong = self.get_argument("susong","")
        self.is_need = self.get_argument("is_need",1)


    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'znys')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir, '风险评估_{}.docx'.format(now),)
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/znys/{}'.format('风险评估_{}.docx'.format(now))
        return result

    def generate_lawapply(self,filepath):
        document = Document()
        #os.path.join(settings.BASE_DIR, 'templates/法律意见书.docx')
        document.styles['Normal'].font.name = u'黑体'  # 可换成word里面任意字体
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置
        run = p.add_run(u"风险评估")
        run.font.size = Pt(28)
        document.add_heading(u'耗时成本', 1)
        document.add_paragraph("申请调解到调解结束平均耗时" + self.mediateTime +"天左右")
        document.add_paragraph("起诉之日到一审判决平均耗时" + self.trialTime +"天左右")
        if int(self.is_need)==0:
            document.add_heading(u'获赔比例', 1)
            document.add_paragraph(u"诉讼获配比例323:100 调解获配比例263:100")
            document.add_heading(u'获赔时间', 1)
            document.add_paragraph(u"诉讼获赔时间：起诉到最终获赔不仅要经过立案、排期、开庭、判决等漫长时间，还会承担证据不足的败诉风险，这就使得自己不但拿不到钱，还付出了打一场官司的成本。")
            document.add_paragraph(u"调解获赔时间：无需等待漫长的法院排期，双方在签完调解协议后可尽快获得钱款。")
        elif int(self.is_need)==1:
            pass
        document.add_heading(u'诉调花费', 1)
        document.add_paragraph("诉讼所花费的金钱：诉讼费" + self.legalFare + "元（以下四种情形，诉讼费减半收取5元：1、以调解方式结案的；2、当事人申请撤诉的；3、当事人拒不到庭，或者中途退庭，视为撤诉的；4、适用简易程序审理的。）")
        document.add_paragraph("调解所花费的金钱：调解费" + self.mediateFare + "元")
        if self.tiaojie:
            document.add_paragraph("调解比例：" + self.tiaojie)
        else:
            pass
        if self.susong:
            document.add_paragraph("诉讼比例：" + self.susong)
        else:
            pass

        document.add_heading(u'信用成本', 1)
        document.add_paragraph("诉讼信用成本：通过起诉解决纠纷，判决书会记载在个人信用记录里。")
        document.add_paragraph("调解信用成本：通过调解解决纠纷，双方达成调解书，不留不良信用记录。")

        # 保存文件
        document.save(filepath)


class PrintpicturenHandler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(PrintpicturenHandler, self).prepare()
        self.type = self.get_argument("type")
        self.materials = self.get_argument("materials")


    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'picture')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir, '诉调流程_{}.docx'.format(now), )
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/picture/{}'.format('诉调流程_{}.docx'.format(now))
        return result

    def generate_lawapply(self, filepath):
        document = Document()
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置
        if int(self.type) == 1:
            run_law_word = p.add_run(u"诉讼流程图")
            run_law_word.font.name = u'宋体'
            run_law_word.font.size = Pt(24)
            f1 = run_law_word._element
            f1.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            pic_path = os.path.join(settings.BASE_DIR, 'templates/susong.png')
            document.add_picture(pic_path, width=Inches(4),height=Inches(5))
        else:
            run_law_word = p.add_run(u"调解流程图")
            run_law_word.font.name = u'宋体'
            run_law_word.font.size = Pt(24)
            f1 = run_law_word._element
            f1.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            pic_path = os.path.join(settings.BASE_DIR, 'templates/tiaojie.png')
            document.add_picture(pic_path, width=Inches(4), height=Inches(5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        law_font = document.add_paragraph()
        run_law = law_font.add_run(u'所需材料：')
        run_law.font.name = u'宋体'
        run_law.font.size = Pt(13)
        r = run_law._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        document.add_paragraph()
        run_materials = law_font.add_run(self.materials)
        run_materials.font.name = u'宋体'
        run_materials.font.size = Pt(13)
        r = run_materials._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        # 保存文件
        document.save(filepath)


class GZLawapplyHandler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(GZLawapplyHandler, self).prepare()

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()
        self.apply_name = self.get_argument("apply_name",'')
        self.apply_mobile = self.get_argument("apply_mobile", '')
        self.apply_degree = self.get_argument('apply_degree', '')
        self.apply_health = self.get_argument('apply_health', '')
        self.apply_company = self.get_argument('apply_company', '')
        self.apply_card = self.get_argument('apply_card', '')
        self.apply_domicile_place = self.get_argument('apply_domicile_place', '')
        self.apply_profession = self.get_argument('apply_profession', '')
        self.agent_name = self.get_argument('agent_name', '')
        self.agent_apply_relationship = self.get_argument('agent_apply_relationship', '')
        self.agent_card = self.get_argument('agent_card_number', '')
        self.agent_mobile = self.get_argument('agent_mobile', '')
        self.agent_address = self.get_argument('agent_address', '')
        self.is_need_declare = self.get_argument('is_need_declare', '')
        self.situation = self.get_argument('situation', '')
        self.description = self.get_argument('description', '')
        self.trialstage_civil = self.get_argument('trialstage_civil', '')
        self.trialstage_administration = self.get_argument('trialstage_administration', '')
        self.trialstage_penal = self.get_argument('trialstage_penal', '')
        self.payment = self.get_argument('payment', '')
        self.receive_address = self.get_argument('receive_address', '')
        self.receive_postal_code = self.get_argument('receive_postal_code', '')
        self.receive_name = self.get_argument('receive_name', '')
        self.receive_mobile = self.get_argument('receive_mobile', '')

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'yitiji')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir,'法律援助申请表广州司法局{}.doc'.format(now))
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/yitiji/{}'.format('法律援助申请表广州司法局{}.doc'.format(now))
        return result

    def generate_lawapply(self, filepath):
        temp_docx = Document(os.path.join(settings.BASE_DIR,'templates/法律援助申请表广州司法局.docx'))
        table = temp_docx.tables[0]
        table.cell(0, 3).text = table.cell(0, 3).text.format(name=self.apply_name)
        table.cell(1, 3).text = table.cell(1, 3).text.format(mobile=self.apply_mobile)
        table.cell(1, 6).text = table.cell(1, 6).text.format(degree=self.apply_degree)
        table.cell(1, 8).text = table.cell(1, 8).text.format(health=self.apply_health)
        table.cell(2, 3).text = table.cell(2, 3).text.format(apply_company=self.apply_company)
        table.cell(0, 7).text = table.cell(0, 7).text.format(apply_card=self.apply_card)
        table.cell(3, 3).text = table.cell(3, 3).text.format(apply_domicile_place=self.apply_domicile_place)
        table.cell(2, 8).text = table.cell(2, 8).text.format(job=self.apply_profession)
        table.cell(4, 0).text = table.cell(4, 0).text.format(agent_name=self.agent_name,
                                                             agent_card=self.agent_card,
                                                             relate=self.agent_apply_relationship,
                                                             agent_mobile=self.agent_mobile,
                                                             agent_address=self.agent_address)

        cell = table.cell(5,5)
        format_text = cell.text
        if self.is_need_declare:
            if self.is_need_declare=="否":
                #for cate in self.agent_type:
                pos = format_text.find("否 (若否，请填写《经济困难申报表》)")
                format_text = "□" + format_text[:pos] + "☑否" + format_text[pos + 1:]
                cell.text = format_text
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(format_text)
                run.font.name = '宋体'
                run.font.size = Pt(11)
            else:
                pos = format_text.find("是（若是，请在以下相应情形打“√”)")
                format_text = format_text[:pos] + "☑是" + format_text[pos + 1:]
                new_str = format_text.replace('否', '□否', 1)
                cell.text = new_str
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(new_str)
                run.font.name = '宋体'
                run.font.size = Pt(11)

        cell_data = table.cell(6, 0)
        format_text_data = cell_data.text
        if self.situation:
            for cate in eval(self.situation):
                pos = format_text_data.find("□ " + cate)
                format_text_data =format_text_data[:pos] + "☑" + format_text_data[pos + 1:]
                cell_data.text = format_text_data
                paragraph = cell_data.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(format_text_data)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)

        table.cell(8, 0).text = table.cell(8, 0).text.format(description=self.description)
        # table.cell(10, 0).text = table.cell(10, 0).text.format(trialstage_civil=self.trialstage_civil)
        item_data = table.cell(10, 0)
        data_text_formart = item_data.text
        if self.trialstage_civil:
            pos = data_text_formart.find("□ ")
            data_text_formart = data_text_formart[:pos] + "☑" + data_text_formart[pos + 1:]
            s1 = list(data_text_formart)
            s1[4] = "□"
            for cs in eval(self.trialstage_civil):
                if  int(cs) == 1:
                    s1[4] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 2:
                    s1[17] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 3:
                    s1[30] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 4:
                    s1[41] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 5:
                    s1[54] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 6:
                    s1[65] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 7:
                    s1[78] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 8:
                    s1[91] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 9:
                    s1[100] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 10:
                    s1[113] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 11:
                    s1[126] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 12:
                    s1[137] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 13:
                    s1[145] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 14:
                    s1[154] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 15:
                    s1[167] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 16:
                    s1[178] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 17:
                    s1[191] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 18:
                    s1[205] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 19:
                    s1[216] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 20:
                    s1[229] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)
                elif int(cs) == 21:
                    s1[242] = "☑"
                    new_data_str = ''.join(s1)
                    item_data.text = data_text_formart
                    paragraph = item_data.paragraphs[0]
                    paragraph.clear()
                    run = paragraph.add_run(new_data_str)
                    run.font.name = '宋体'
                    run.font.size = Pt(10.5)


        table.cell(12, 2).text = table.cell(12, 2).text.format(
                                                               receive_name=self.receive_name,
                                                               receive_address=self.receive_address,
                                                               receive_mobile=self.receive_mobile,
                                                               receive_postal_code=self.receive_postal_code)
        cell_chose = table.cell(12, 2)
        format_chose_data = cell_chose.text
        if self.payment:
            if self.payment == '1':
                pos = format_chose_data.find("□ " + "自行前往法律援助机构领取")
                format_chose_data = format_chose_data[:pos] + "☑" + format_chose_data[pos + 1:]
                cell_chose.text = format_chose_data
                paragraph = cell_chose.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(format_chose_data)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
            else:
                pos = format_chose_data.find("□ " +"邮寄送达")
                format_chose_data = format_chose_data[:pos] + "☑" + format_chose_data[pos + 1:]
                cell_chose.text = format_chose_data
                paragraph = cell_chose.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(format_chose_data)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
        temp_docx.save(filepath)


class PrintIndictmentHandler(handlers.NoAuthHandler):
    """
@name 打印起诉状
@description 打印起诉状

@path
  /intelligentpretrial/docdispose/indictment:
    post:
      tags:
      - "打印起诉状"
      parameters:
      - name: "yg"
        in: "body"
        description: "原告信息"
        required: true
        schema:
          $ref: "#/definitions/Yg"
      - name: "ygdlr"
        in: "body"
        description: "原告代理人信息"
        required: true
        schema:
          $ref: "#/definitions/Dlr"
      - name: "bg"
        in: "body"
        description: "被告人信息"
        schema:
          $ref: "#/definitions/Yg"
        required: true
      - name: "bgdlr"
        in: "body"
        description: "被告代理人信息"
        required: true
        schema:
          $ref: "#/definitions/Dlr"
      - name: "zz"
        in: "query"
        description: "当事人主张"
        type: "string"
        required: true
      - name: "ssly"
        in: "query"
        description: "当事人事实和理由"
        type: "string"
        required: true

      responses:
        200:
          description: "返回成功"
          schema:
            $ref: "#/definitions/QWata"
        400:
          description: "参数错误"
@endpath

@definitions
  QWata:
    type: "object"
    properties:
      data:
        $ref: "#/definitions/OPath"
      msg:
        type: "string"
      code:
        type: "integer"
        description: "200返回成功，400返回失败"
        enum:
        - "200"
        - "400"
  OPath:
    type: "object"
    properties:
      filepath:
        type: "string"
        description: "起诉状文件路径"
  Yg:
    type: "object"
    properties:
      name:
        type: "string"
        description: "名字"
      sex:
        type: "string"
        description: "性别"
      id:
        type: "string"
        description: "身份证号码"
      address:
        type: "string"
        description: "住址"
      phone:
        type: "string"
        description: "联系电话"
      frzw:
        type: "string"
        description: "法人职务"
      frname:
        type: "string"
        description: "法人名字"
      fraddress:
        type: "string"
        description: "法人住址"
  Dlr:
    type: "object"
    properties:
      type:
        type: "string"
        description: "类型"
      name:
        type: "string"
        description: "名字"
      sex:
        type: "string"
        description: "性别"
      work:
        type: "string"
        description: "工作"
      address:
        type: "string"
        description: "地址"
      lsname:
        type: "string"
        description: "律所名称"
      lsaddress:
        type: "string"
        description: "律所地址"
      lsphone:
        type: "string"
        description: "律所电话"
@enddefinitions
    """
    @decorator.handler_except
    def parse_param(self):
        super(PrintIndictmentHandler, self).prepare()
        self.yg = demjson.decode(self.get_argument("yg"))
        self.yg_dlr = demjson.decode(self.get_argument("ygdlr"))
        self.beigao = demjson.decode(self.get_argument("bg"))
        self.beigao_dl = demjson.decode(self.get_argument("bgdlr"))
        self.zhuzhang = self.get_argument('zz', '')
        self.ssly = self.get_argument('ssly', '')
        self.qid = self.get_argument('qid', '')
        self.court= self.get_argument('court', '')

        '''原告信息'''
        self.yuangao = []
        self.ygtype = self.yg.setdefault('type','')
        self.yuangao.append(self.yg.setdefault('name',''))
        self.yuangao.append(self.yg.setdefault('sex',''))
        self.yuangao.append(self.yg.setdefault('id',''))
        self.yuangao.append(self.yg.setdefault('family',''))
        self.yuangao.append(self.yg.setdefault('address',''))
        self.yuangao.append(self.yg.setdefault('birthday',''))
        self.yuangao.append(self.yg.setdefault('phone',''))
        self.yuangao.append(self.yg.setdefault('frname',''))
        self.yuangao.append(self.yg.setdefault('fraddress',''))
        self.yuangao.append(self.yg.setdefault('frzw',''))
        '''原告代理人'''
        self.ygdlr = []
        del self.yg_dlr["type"]
        for g_dl,y_dl in self.yg_dlr.items():
            if y_dl != "":
                self.ygdlr.append(y_dl)
        '''被告'''
        self.bg = []
        self.bgtype = self.beigao.setdefault('type', '')
        self.bg.append(self.beigao.setdefault('name', ''))
        self.bg.append(self.beigao.setdefault('sex', ''))
        self.bg.append(self.beigao.setdefault('id', ''))
        self.bg.append(self.beigao.setdefault('address', ''))
        self.bg.append(self.beigao.setdefault('birthday', ''))
        self.bg.append(self.beigao.setdefault('phone', ''))
        self.bg.append(self.beigao.setdefault('frname', ''))
        self.bg.append(self.beigao.setdefault('fraddress', ''))
        self.bg.append(self.beigao.setdefault('frzw', ''))
        '''被告代理人'''
        self.bgdlr = []
        del self.beigao_dl["type"]
        for k_dl,v_dl in self.beigao_dl.items():
            if v_dl != "":
                self.bgdlr.append(v_dl)

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'indictment')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir, '起诉状_{}.docx'.format(now), )
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/indictment/{}'.format('起诉状_{}.docx'.format(now))
        return result

    def generate_lawapply(self, filepath):
        document = Document()
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置
        run_law_word = p.add_run(u"民事起诉状")
        run_law_word.font.name = u'宋体'
        run_law_word.font.size = Pt(24)
        gs=""
        for yg in self.yuangao:
            if yg != "":
                gs = gs + yg + '、'
        gs = gs[:len(gs) - 1]
        paragraph=document.add_paragraph()
        run_yg = paragraph.add_run("         原告:" + gs)
        run_yg.font.size = Pt(14)
        gq = ""
        print("self.ygdlr",self.ygdlr)
        if self.ygdlr ==[]:
            pass
        else:
            for ygdl in self.ygdlr:
                if ygdl != "":
                    gq = gq + ygdl + '、'
            gq = gq[:len(gq) - 1]

            yg_wtdlr = document.add_paragraph()
            run = yg_wtdlr.add_run("         " + "委托代理人:" + gq)
            run.font.size = Pt(14)

        gw = ""
        for ber in self.bg:
            if ber != "":
                gw = gw + ber + '、'
        gw = gw[:len(gw) - 1]
        beigao=document.add_paragraph()
        run_bg = beigao.add_run("         被告:" + gw)
        run_bg.font.size = Pt(14)
        ge = ""
        if self.bgdlr==[]:
            pass
        else:
            for bgdl in self.bgdlr:
                if bgdl != "":
                    ge = ge + bgdl + '、'
            ge = ge[:len(ge) - 1]
        bg_wtdlr = document.add_paragraph()
        run_dl = bg_wtdlr.add_run("         " + "委托代理人:" + ge)
        run_dl.font.size = Pt(14)

        ssqq = document.add_paragraph()
        run_ssqq = ssqq.add_run("诉讼请求:")
        run_ssqq.font.size = Pt(14)

        zz_data = self.zhuzhang.split('；')
        for zz in zz_data:
            zz_p=document.add_paragraph()
            run_zz = zz_p.add_run("         "+zz)
            run_zz.font.size = Pt(14)

        sslybt = document.add_paragraph()
        run_ss1 = sslybt.add_run("事实和理由:")
        run_ss1.font.size = Pt(14)

        ssyly=document.add_paragraph()
        run_ss = ssyly.add_run("         "+self.ssly)
        run_ss.font.size = Pt(14)
        rmfy=document.add_paragraph()
        if self.court !="":
            run_rmfy = rmfy.add_run(self.court)
            run_rmfy.font.size = Pt(14)
        else:
            run_rmfy = rmfy.add_run("____人民法院")
            run_rmfy.font.size = Pt(14)
        jzr=document.add_paragraph()
        run_jzr = jzr.add_run("\t\t\t\t\t\t\t具状人：" +self.yuangao[0])
        run_jzr.font.size = Pt(14)
        from datetime import datetime
        x = datetime.now()
        y = datetime.strftime(x, "%Y年%m月%d日")
        rqsj=document.add_paragraph()
        run_rq = rqsj.add_run("\t\t\t\t\t\t\t日期：" + y)
        run_rq.font.size = Pt(14)
        # 保存文件
        document.save(filepath)


class RiskDownloadHandler(handlers.NoAuthHandler):
    """
@name 下载风险评估
@description 下载风险评估

@path
  /intelligentpretrial/docdispose/downassess:
    post:
      tags:
      - "下载风险评估"
      parameters:
      - name: "demand"
        in: "query"
        description: "诉讼获赔比例"
        type: "string"
        required: true
      - name: "compensat"
        in: "query"
        description: "调解获配比例"
        type: "string"
        required: true
      - name: "mediateTime"
        in: "query"
        description: "申请调解到调解结束时间"
        type: "string"
        required: true
      - name: "trialTime"
        in: "query"
        description: "起诉之日到判决时间"
        required: true
        type: "string"
      - name: "susongfei"
        in: "query"
        description: "诉讼费内容"
        type: "string"
        required: true
      - name: "login"
        in: "query"
        type: "string"
        description: "1为登录版本，0为非登录版本"
        required: false
      responses:
        200:
          description: "返回成功"
          schema:
            $ref: "#/definitions/Data_msg"
        400:
          description: "参数错误"
@endpath

@definitions
  Data_msg:
    type: "object"
    properties:
      data:
        $ref: "#/definitions/filepath"
      msg:
        type: "string"
      code:
        type: "integer"
        description: "200返回成功，400返回失败"
        enum:
        - "200"
        - "400"
  filepath:
    type: "object"
    properties:
      filepath:
        type: "string"
        description: "返回的下载地址（路径）"
@enddefinitions
    """
    @decorator.handler_except
    def parse_param(self):
        super(RiskDownloadHandler, self).prepare()
        self.demand = self.get_argument("demand","")#索要
        self.compensat = self.get_argument("compensat","")#获赔
        self.mediateTime = self.get_argument("mediateTime","")#调解时间
        self.trialTime = self.get_argument("trialTime","")#诉讼时间
        self.susongfei = self.get_argument("susongfei","")#诉讼费
        self.login = self.get_argument("login", 0)

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'znys')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir, '风险评估_{}.docx'.format(now),)
        self.generate_lawapply(filepath)
        result['data'] = {}
        if int(self.login) == 0:
            domain ='http://noauth.aegis-info.com'
        else:
            domain = 'http://intelligentpretrial.aegis-info.com'
        result['data']['filepath'] = domain+'/static/znys/{}'.format('风险评估_{}.docx'.format(now))
        return result

    def generate_lawapply(self,filepath):
        document = Document()
        #os.path.join(settings.BASE_DIR, 'templates/法律意见书.docx')
        document.styles['Normal'].font.name = u'黑体'  # 可换成word里面任意字体
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落文字居中设置
        run = p.add_run(u"风险评估")
        run.font.size = Pt(28)
        document.add_heading(u'时间成本', 1)
        document.add_paragraph("申请调解到调解结束平均耗时" + self.mediateTime +"天左右")
        document.add_paragraph("起诉之日到一审判决平均耗时" + self.trialTime +"天左右")
        if self.demand:
            document.add_heading(u'获赔比例', 1)
            document.add_paragraph(u"诉讼获赔比例："+self.demand+"调解获赔比例："+self.compensat)
        document.add_heading(u'获赔时间', 1)
        document.add_paragraph(u"诉讼获赔时间：起诉到最终获赔不仅要经过立案、排期、开庭、判决等漫长时间，还会承担证据不足的败诉风险，这就使得自己不但拿不到钱，还付出了打一场官司的成本。")
        document.add_paragraph(u"调解获赔时间：无需等待漫长的法院排期，双方在签完调解协议后可尽快获得钱款。")
        document.add_heading(u'金钱成本', 1)
        if self.susongfei:
            document.add_paragraph("诉讼费："+self.susongfei)
        else:
            pass
        document.add_heading(u'其他风险', 1)
        document.add_paragraph(u"诉讼处处有风险，起诉不符合条件可能不会受理、诉讼请求不适当可能不会审理、"
                               "超过诉讼时效诉讼请求可能不被支持、授权不明代理人发表的意见可能不具有法律"
                               "效力、不按时交纳诉讼费用可能被撤诉、不提供或不充分提供证据可能承担败诉的后果，"
                               "既使胜诉了，若败诉的人无财产可供执行，还需承担执行不能的风险。")
        document.add_paragraph(u"重要提示： 结果仅供参考，具体案件以法院审理为准。")
        # 保存文件
        document.save(filepath)


class BusinessPrintHandler(handlers.NoAuthHandler):
    @decorator.handler_except
    def parse_param(self):
        super(BusinessPrintHandler, self).prepare()

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):
        result = self.init_response_data()
        self.name = self.get_argument("name", '')#当事人名称
        self.sex = self.get_argument("sex", '')#当事人性别
        self.idnumber = self.get_argument("idnumber", '')#身份证号
        self.age = self.get_argument("age", '')#年龄
        self.type = self.get_argument("type", '')

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'business')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir, self.type+'_{}.docx'.format(now))
        self.business_print(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/business/{}'.format(self.type+'_{}.docx'.format(now))
        return result
    def business_print(self, filepath):
        temp_path=""
        if self.type=="人民调解":
            temp_path='templates/人民调解业务.docx'
        elif self.type=="仲裁服务":
            temp_path = 'templates/仲裁服务业务.docx'
        elif self.type=="公证服务":
            temp_path = 'templates/公证服务业务.docx'
        elif self.type=="司法鉴定":
            temp_path = 'templates/司法鉴定业务.docx'
        elif self.type=="法律咨询":
            temp_path = 'templates/法律咨询业务.docx'
        elif self.type=="法律援助":
            temp_path = 'templates/法律援助业务.docx'
        elif self.type=="行政复议":
            temp_path = 'templates/行政复议业务.docx'
        temp_docx = Document(os.path.join(settings.BASE_DIR,temp_path))
        table = temp_docx.tables[0]
        if self.type == "公证服务":
            table.cell(0, 2).text = table.cell(0, 2).text.format(name=self.name)
            table.cell(0, 4).text = table.cell(0, 4).text.format(sex=self.sex)
            table.cell(0, 6).text = table.cell(0, 6).text.format(age=self.age)
            table.cell(2, 2).text = table.cell(2, 2).text.format(idnumber=self.idnumber)
            cell = table.cell(0, 6)
            format_text = cell.text
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '华文细黑'
            run.font.size = Pt(11.5)
            cell = table.cell(2, 2)
            format_text = cell.text
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '宋体'
            run.font.size = Pt(11.5)
            temp_docx.save(filepath)
        else:
            table.cell(0, 1).text = table.cell(0, 1).text.format(name=self.name)
            table.cell(0, 3).text = table.cell(0, 3).text.format(sex=self.sex)
            table.cell(0, 5).text = table.cell(0, 5).text.format(age=self.age)
            table.cell(2, 1).text = table.cell(2, 1).text.format(idnumber=self.idnumber)
            cell = table.cell(0, 5)
            format_text = cell.text
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '华文细黑'
            run.font.size = Pt(11.5)
            cell = table.cell(2, 1)
            format_text = cell.text
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '宋体'
            run.font.size = Pt(11.5)
            temp_docx.save(filepath)

handlers = [
    (r"/intelligentpretrial/docdispose/lawopinion", LawopinionHandler),
    (r"/intelligentpretrial/docdispose/riskassessment",RiskassessmentHandler),
    (r"/intelligentpretrial/docdispose/printpic",PrintpicturenHandler),
    (r"/intelligentpretrial/docdispose/gzlawapply",GZLawapplyHandler),
    (r"/intelligentpretrial/docdispose/indictment",PrintIndictmentHandler),
    (r"/intelligentpretrial/docdispose/downassess",RiskDownloadHandler),
    (r"/intelligentpretrial/docdispose/business",BusinessPrintHandler)
]
