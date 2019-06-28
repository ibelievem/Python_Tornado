# @Time    : 2018/7/13
# @Author  : Dongrj


import datetime
import demjson
import json
import os
import re
import time

from docx import Document
from docx.shared import Pt
from microserver.core.exceptions import ArgumentTypeError
from microserver.core.handlers import handlers
from microserver.conf import settings
from microserver.utils import decorator
from tornado import gen

from .personal_damage_compensation_calculator.calculators import *
from calculator.models import ZNYSModel
from engine.collections import Params


class CalculatorHandler(handlers.BaseHandler):
    """
@name 未签劳动合同计算器
@description 未签劳动合同计算器

@path
  /intelligentpretrial/android/law_push/Unsigned_Contract_Handler:
    get:
      tags:
      - "未签劳动合同计算器"
      parameters:
      - name: "start"
        in: "query"
        description: "入职时间"
        type: "string"
        required: true
      - name: "end"
        in: "query"
        description: "截止时间"
        type: "string"
        required: true
      - name: "salary"
        in: "query"
        description: "月平均收入"
        type: "string"
        required: true
      - name: "time_of_award"
        in: "query"
        description: "劳动仲裁时间"
        type: "string"
        required: true
      - name: "qid"
        in: "query"
        description: "选项记录的id"
        type: "string"
        required: true
      responses:
        200:
          description: "返回成功"
          schema:
            $ref: "#/definitions/Calculator"
        400:
          description: "返回失败"
@endpath

@definitions
  Calculator:
    type: "object"
    properties:
      data:
        $ref: "#/definitions/Lawjisu"
      msg:
        type: "string"
      code:
        type: "integer"
        description: "200返回成功，400返回失败"
        enum:
        - "200"
        - "400"
  Lawjisu:
    type: "object"
    properties:
      money:
        type: "integer"
        description: "双倍工资赔偿金额"
      answer:
        type: "string"
        description: "由计算器得出的文案"
@enddefinitions
    """
    @decorator.handler_except
    def prepare(self):
        """在请求方法 get、post等执行前调用，进行通用的参数初始化，支持协程"""
        self.start = self.get_argument('start', '1')
        self.end = self.get_argument("end", '1')
        self.salary = int(self.get_argument("salary", '0'))
        self.time_of_award = self.get_argument("time_of_award", '1') or '1'
        self.qid = self.get_argument("qid", '')

    @decorator.future_except
    @gen.coroutine
    def get(self, *args, **kwargs):
        """IO操作"""
        result = self.init_response_data()
        result['data'] = {}
        money, answer = self.arbitration_costs(self.start, self.end, self.salary, self.time_of_award, self.qid)
        result['data']['money'] = money
        result['data']['answer'] = answer
        self.finish(result)

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):
        """耗时操作"""
        result = self.init_response_data()
        return result

    def arbitration_costs(self, start, end, salary, time_of_award, qid):
        znysmodel = ZNYSModel()
        time_s = start.split('-')
        time_e = end.split('-')
        tmp_time_of_award = time_of_award
        leap_month = [31, 29, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
        no_leap_month = [31, 28, 31, 30, 31, 30, 31, 31, 30, 31, 30, 31]
        if time_of_award == '1' or time_of_award == None:
            time_of_award = datetime.datetime.now().strftime("%Y-%m-%d")
        time_of_award = time_of_award.split('-')
        int_s = [int(ii) for ii in time_s]
        int_e = [int(ii) for ii in time_e]
        int_award = [int(ii) for ii in time_of_award]
        data_s = datetime.date(int_s[0], int_s[1], int_s[2])
        data_e = datetime.date(int_e[0], int_e[1], int_e[2])
        data_award = datetime.date(int_award[0], int_award[1], int_award[2])
        delta = datetime.timedelta(days=30)
        delta1 = datetime.timedelta(days=1)
        delta365 = datetime.timedelta(days=365)
        if (int_s[0] % 400 == 0) or (int_s[0] % 4 == 0 and int_s[0] % 100 != 0):
            delta_n = datetime.timedelta(days=leap_month[int_s[1] - 1])
        else:
            delta_n = datetime.timedelta(days=no_leap_month[int_s[1] - 1])
        data_s_30 = data_s + delta_n
        data_e = datetime.date(int_e[0], int_e[1], int_e[2])
        data_award = datetime.date(int_award[0], int_award[1], int_award[2])
        delta = datetime.timedelta(days=30)
        delta1 = datetime.timedelta(days=1)
        delta365 = datetime.timedelta(days=365)
        if (int_s[0] == int_e[0] and ((int_s[1] + 1 == int_e[1] and int_s[2] < int_e[2]) or int_s[1] == int_e[1])) or \
                (int_s[0] + 1 == int_e[0] and int_s[1] == 12 and int_e[1] == 1 and int_s[2] < int_e[2]):
            money = 0
            answer = "    你是" + time_s[0] + "年" + time_s[1] + "月" + time_s[2] + "日到单位工作，截止" + time_e[0] + "年" + time_e[
                1] + "月" + time_e[2] + "日你还未签劳动合同且在该单位工作不满一个月。" \
                "按照《劳动合同法》第10条的规定，单位应当在你入职的一个月内与你签订劳动合同，所以单位未与你签劳动合同仍在合法期限内，" \
                "而请求单位支付未签劳动合同的双倍工资是从工作满一个月的次日开始起算的，鉴于你工作不满一个月，所以无法请求单位支付未签订劳动合同的双倍工资。"
        elif int_s[0] == int_e[0] or (int_s[0] + 1 == int_e[0] and int_s[1] > int_e[1]) \
                or (int_s[0] + 1 == int_e[0] and int_s[1] == int_e[1] and int_s[2] > int_e[2]):
            month = int_e[1] - int_s[1]
            day = int_e[2] - int_s[2]
            if month > 0:
                money = round(salary * (month - 1) + salary * day / 30, 2)
            else:
                money = round(salary * (month + 11) + salary * day / 30, 2)
            data_ee = data_e + delta1
            data_eee = data_e + delta365
            answer = "    你是" + time_s[0] + "年" + time_s[1] + "月" + time_s[2] + "日到单位工作，截止" + time_e[0] + "年" + time_e[
                1] + "月" + time_e[2] + "日你还未签劳动合同且在该单位工作不满一年。 " \
                "按照《劳动合同法》第10条的规定，单位应当在你入职的一个月内与你签订劳动合同。但你工作满一个月了，单位也未和你签劳动合同，因此单位已经超过了签订劳动合同的法定时间。" \
                "所以你可以根据《劳动合同法实施条例》第6条向公司主张<b>" + str(data_s_30.year) + "年" + str(
                data_s_30.month) + "月" + str(data_s_30.day) + "日至" + time_e[0] + "年" + time_e[1] + "月" + time_e[2] + "日</b>这段期间的双倍工资。\n" \
                "    但是根据《劳动争议调解仲裁法》第27条的规定，你需要在权利受到侵害之日起1年内主张你的两倍工资才有可能得到支持。" \
                "也就是说你需要在<b>" + str(data_eee.year) + "年" + str(data_eee.month) + "月" + str(data_eee.day) + "日</b>之前向劳动仲裁委员会申请仲裁，一旦超过此时间段就难以得到支持。" \
                "由于你在该单位工作不满一年，在1年有效仲裁时效期间内主张未签劳动合同两倍工资差额，则从工作满一个月之次日起到未签劳动合同之日期间主张两倍工资差额可以得到认可，" \
                "即可主张" + str(data_s_30.year) + "年" + str(data_s_30.month) + "月" + str(data_s_30.day) + "日至" + time_e[0] + "年" + time_e[1] + "月" + time_e[2] + "日期间未签劳动合同的两倍工资差额" + str(money)+ "元。"
        elif (data_e - data_s).days >= 365 and (data_award - data_s).days <= 365*2:#and int_s[2] <= int_e[2]
            data_award_365 = data_award - delta365
            data_award_3651 = data_award - delta365 - delta1
            data_s_365 = data_s + delta365 - delta1
            data_s_3651 = data_s + delta365
            data_s_3652 = data_s + delta365 + delta365
            if (data_award - data_s).days <= 365:
                month = data_award.month - data_s.month
                day = data_award.day - data_s.day
                if month > 0:
                    money = round(salary * (month - 1) + salary * day / 30, 2)
                else:
                    money = round(salary * (month + 11) + salary * day / 30, 2)
            else:
                month = data_s_365.month - data_award.month
                day = data_s_365.day - data_award.day
                if month > 0:
                    money = round(salary * month + salary * day / 30, 2)
                else:
                    money = round(salary * (month + 12) + salary * day / 30, 2)
            if tmp_time_of_award == '1' or tmp_time_of_award == None:
                answer = "    你是" + time_s[0] + "年" + time_s[1] + "月" + time_s[2] + "日到单位工作，截止" + time_e[0] + \
                         "年" + time_e[1] + "月" + time_e[2] + "日你还未签劳动合同。由于未明确申请劳动仲裁时间，因此你" \
                        "最迟需要在" + str(data_s_3652.year) + "年" + str(data_s_3652.month) + "月" + str(data_s_3652.day) + \
                         "日前申请仲裁才有可能拿到部分的双倍工资。"
            else:
                answer = "    你是" + time_s[0] + "年" + time_s[1] + "月" + time_s[2] + "日到单位工作，截止" + time_e[0] + "年" + time_e[
                    1] + "月" + time_e[2] + "日你还未签劳动合同且在该单位工作已满一年。" \
                    "按照《劳动合同法》第10条的规定，单位应当在你入职的一个月内与你签订劳动合同。但你工作满一个月了，单位也未和你签劳动合同，因此单位已经超过了签订劳动合同的法定时间。" \
                    "所以你可以根据《劳动合同法实施条例》第7条向公司主张<b>" + str(data_award_365.year) + "年" + str(data_award_365.month) + "月" + str(data_award_365.day) + \
                    "日至" + str(data_s_3651.year) + "年" + str(data_s_3651.month) + "月" + str(data_s_3651.day) + "日</b>这段期间的双倍工资。" \
                    "另外，自用工之日起满一年未签劳动合同，还视为单位自" + str(data_s_3651.year) + "年" + str(data_s_3651.month) + "月" + str(data_s_3651.day) + "日起就与你订立了无固定期限劳动合同，" \
                    "若在工作中没有重大违法违纪行为，而单位将你辞退，则可以要求单位支付经济赔偿金。\n    但是根据《劳动争议调解仲裁法》第27条的规定，你需要在权利受到侵害之日起1年内主张你的两倍工资才有可能得到支持。" \
                    "也就是说你需要在<b>" + str(data_s_3652.year) + "年" + str(data_s_3652.month) + "月" + str(data_s_3652.day) + \
                    "日</b>之前向劳动仲裁委员会申请仲裁，一旦超过此时间段就难以得到支持。由于你在" + str(data_award.year) + "年" + str(data_award.month) + "月" + str(data_award.day) + "日提出仲裁申请，" \
                    "则要求" + str(data_award_3651.year) + "年" + str(data_award_3651.month) + "月" + str(data_award_3651.day) + "日前的未签订劳动合同的两倍工资差额，明显已超过申请仲裁时效期间，" \
                    "故对主张的" + str(data_award_3651.year) + "年" + str(data_award_3651.month) + "月" + str(data_award_3651.day) + "日前的两倍工资差额不予认可，" \
                    "仅支持" + str(data_award_365.year) + "年" + str(data_award_365.month) + "月" + str(data_award_365.day) + "日至" + str(data_s_3651.year) + "年" + str(data_s_3651.month) + "月" + str(data_s_3651.day) + \
                    "日期间未签劳动合同的两倍工资差差额" + str(money)+ "元。"
        else:
            money = 0
            answer = "    你是" + time_s[0] + "年" + time_s[1] + "月" + time_s[2] + "日到单位工作，截止" + time_e[0] + "年" + time_e[1] + "月" + time_e[2] + "日你还未签劳动合同。" \
                    "但是由于你超过了申请劳动仲裁的时效期间，你的请求将不受法律保护，所以你请求单位支付未签劳动合同双倍工资赔偿的主张不会得到支持。"

        body = dict(
            start=start,
            end=end,
            salary=salary,
            time_of_award=time_of_award,
            money=money,
        )
        aa = znysmodel.update_record(body, qid)
        print(aa)
        return money, answer


class LawapplyHandler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(LawapplyHandler, self).prepare()

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()
        print("############################", self.get_argument("word"))
        self.law_apply = demjson.decode(self.get_argument("word"))
        self.apply_name = self.law_apply.setdefault('apply_name', '')
        self.apply_sex = self.law_apply.setdefault('apply_sex', '')
        self.apply_education = self.law_apply.setdefault('apply_education', '')
        self.apply_crowd = self.law_apply.setdefault('apply_crowd', '')
        self.cause_action = self.law_apply.setdefault('cause_action', '')
        self.apply_race = self.law_apply.setdefault('apply_race', '')
        self.apply_card = self.law_apply.setdefault('apply_card', '')
        self.apply_domicile_place = self.law_apply.setdefault('apply_domicile_place', '')
        self.apply_address = self.law_apply.setdefault('apply_address', '')
        self.apply_postal_code = self.law_apply.setdefault('apply_postal_code', '')
        self.apply_mobile = self.law_apply.setdefault('apply_mobile', '')
        self.apply_workspace = self.law_apply.setdefault('apply_workspace', '')
        self.agent_name = self.law_apply.setdefault('agent_name', '')
        self.agent_phone = self.law_apply.setdefault('agent_phone', '')
        self.agent_type = self.law_apply.setdefault('agent_type', '')
        self.description = self.law_apply.setdefault('description', '')
        self.request_items = self.law_apply.setdefault('request_items', '')

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'yitiji')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir,'法律援助申请表_{}.docx'.format(now))
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/yitiji/{}'.format('法律援助申请表_{}.docx'.format(now))
        return result


    def generate_lawapply(self, filepath):
        temp_docx = Document(os.path.join(settings.BASE_DIR,'templates/法律援助申请表.docx'))
        table = temp_docx.tables[0]
        # print(table.cell(2, 13).text,table.cell(2, 14).text,table.cell(2, 15).text,table.cell(2, 16).text)
        table.cell(1, 1).text = table.cell(1, 1).text.format(apply_name=self.apply_name) #申请人姓名
        table.cell(1, 16).text = table.cell(1, 16).text.format(apply_sex=self.apply_sex) #性别
        table.cell(1, 7).text = table.cell(1, 7).text.format(apply_race=self.apply_race) #民族
        table.cell(1, 20).text = table.cell(1, 20).text.format(apply_education=self.apply_education) #学历
        table.cell(2, 2).text = table.cell(2, 2).text.format(apply_crowd=self.apply_crowd)  # 人群类别
        table.cell(2, 13).text = table.cell(2, 13).text.format(cause_action=self.cause_action)  # 案由
        table.cell(4, 2).text = table.cell(4, 2).text.format(apply_domicile_place=self.apply_domicile_place) #户籍所在地
        table.cell(5, 2).text = table.cell(5, 2).text.format(apply_address=self.apply_address) #现住所
        table.cell(6, 6).text = table.cell(6, 6).text.format(apply_postal_code=self.apply_postal_code) #邮政编码
        table.cell(6, 18).text = table.cell(6, 18).text.format(apply_mobile=self.apply_mobile)  # 联系电话
        table.cell(7, 2).text = table.cell(7, 2).text.format(apply_workspace=self.apply_workspace) #工作单位
        table.cell(8, 2).text = table.cell(8, 2).text.format(agent_name=self.agent_name) #代理人姓名
        table.cell(8, 11).text = table.cell(8, 11).text.format(agent_phone=self.agent_phone)  # 代理人电话
        table.cell(11, 0).text = table.cell(11, 0).text.format(description=self.description,request_items=self.request_items) #描述and请求事项和请求数额
        if len(self.apply_card) == 18:
            table.cell(3, 2).text = self.apply_card[0] #1
            table.cell(3, 3).text = self.apply_card[1]#2
            table.cell(3, 4).text = self.apply_card[2]#3
            table.cell(3, 5).text = self.apply_card[3]#4
            table.cell(3, 6).text = self.apply_card[4]#5

            table.cell(3, 8).text = self.apply_card[5]#6
            table.cell(3, 9).text = self.apply_card[6]#7
            table.cell(3, 10).text = self.apply_card[7]#8
            table.cell(3, 12).text = self.apply_card[8]#9
            table.cell(3, 15).text = self.apply_card[9]#10
            table.cell(3, 16).text = self.apply_card[10]#11

            table.cell(3, 17).text = self.apply_card[11]#12
            table.cell(3, 18).text = self.apply_card[12]#13
            table.cell(3, 19).text = self.apply_card[13]#14
            table.cell(3, 20).text = self.apply_card[14]#15
            table.cell(3, 21).text = self.apply_card[15]#16
            table.cell(3, 22).text = self.apply_card[16]#17
            table.cell(3, 23).text = self.apply_card[17]#18
        else:
            raise Exception("身份证号码格式错误")
        cell = table.cell(8,22)
        format_text = cell.text
        if self.agent_type =='法定代理人':
            #for cate in self.agent_type:
            pos = format_text.find("□" + "法定")
            format_text = format_text[:pos] + "√" + format_text[pos + 1:]
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
        elif self.agent_type =='委托代理人':
            pos = format_text.find("□" + "委托")
            format_text = format_text[:pos] + "√" + format_text[pos + 1:]
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)

        temp_docx.save(filepath)


class ZNYS_Labor_Handler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(ZNYS_Labor_Handler, self).prepare()

    @decorator.threadpoll_executor
    def get(self, *args, **kwargs):
        result = self.init_response_data()
        content = self.get_argument("content", '')
        result['data'] = self.is_ZNYS_Labor(content)
        return result
    def is_ZNYS_Labor(self,content):
        not_word_list = ["没", "未", "无", "不"]
        not_word = "|".join(not_word_list)
        company_list = ["公司", "单位", "老板", "企业"]
        company_word = "|".join(company_list)
        work_list = ["上班", "工作", "打工", "做事", "干活"]
        work_word = "|".join(work_list)
        double_list = ["两倍", "双倍", "二倍"]
        double_word = "|".join(double_list)
        match_tag_regex_labor = "(.*(" + not_word + ").*(" + company_word + ").*合同.*赔偿.*)|" \
                                                                            "(.*(" + not_word + ").*劳动合同.*赔偿.*)|" \
                                                                                                "(.*(" + work_word + ").*(" + not_word + ").*合同.*)|" \
                                                                                                                                         "(.*(" + not_word + ").*合同.*(" + double_word + ").*工资.*)|" \
                                                                                                                                                                                        "(.*(" + company_word + ").*(" + not_word + ").*合同.*赔偿.*)"
        if re.match(match_tag_regex_labor, content):
            return 9462
        else:
            return 0


class LawAidFormHandler(handlers.NoAuthHandler):

    @decorator.handler_except
    def parse_param(self):
        super(LawAidFormHandler, self).prepare()

    @decorator.threadpoll_executor
    def post(self, *args, **kwargs):

        self.parse_param()
        result = self.init_response_data()
        print("############################", self.get_argument("word"))
        self.law_apply = demjson.decode(self.get_argument("word"))
        self.apply_name = self.law_apply.setdefault('apply_name', '')
        self.apply_sex = self.law_apply.setdefault('apply_sex', '')
        self.apply_education = self.law_apply.setdefault('apply_education', '')
        self.apply_crowd = self.law_apply.setdefault('apply_crowd', '')
        self.cause_action = self.law_apply.setdefault('cause_action', '')
        self.apply_race = self.law_apply.setdefault('apply_race', '')
        self.apply_card = self.law_apply.setdefault('apply_card', '')
        self.apply_domicile_place = self.law_apply.setdefault('apply_domicile_place', '')
        self.apply_address = self.law_apply.setdefault('apply_address', '')
        self.apply_postal_code = self.law_apply.setdefault('apply_postal_code', '')
        self.apply_mobile = self.law_apply.setdefault('apply_mobile', '')
        self.apply_workspace = self.law_apply.setdefault('apply_workspace', '')
        self.agent_name = self.law_apply.setdefault('agent_name', '')
        self.agent_phone = self.law_apply.setdefault('agent_phone', '')
        self.agent_type = self.law_apply.setdefault('agent_type', '')
        self.description = self.law_apply.setdefault('description', '')
        self.request_items = self.law_apply.setdefault('request_items', '')

        now = str(int(time.time()))
        filedir = os.path.join(settings.get_static_absolute_path(), 'yitiji')
        if not os.path.exists(filedir):
            os.makedirs(filedir)
        filepath = os.path.join(filedir,'法律援助申请表general_{}.docx'.format(now))
        self.generate_lawapply(filepath)
        result['data'] = {}
        result['data']['filepath'] = '/static/yitiji/{}'.format('法律援助申请表general_{}.docx'.format(now))
        return result


    def generate_lawapply(self, filepath):
        temp_docx = Document(os.path.join(settings.BASE_DIR,'templates/法律援助申请表general.docx'))
        table = temp_docx.tables[0]
        # print(table.cell(2, 13).text,table.cell(2, 14).text,table.cell(2, 15).text,table.cell(2, 16).text)
        table.cell(1, 1).text = table.cell(1, 1).text.format(apply_name=self.apply_name) #申请人姓名
        table.cell(1, 16).text = table.cell(1, 16).text.format(apply_sex=self.apply_sex) #性别
        table.cell(1, 7).text = table.cell(1, 7).text.format(apply_race=self.apply_race) #民族
        table.cell(1, 20).text = table.cell(1, 20).text.format(apply_education=self.apply_education) #学历
        table.cell(2, 2).text = table.cell(2, 2).text.format(apply_crowd=self.apply_crowd)  # 人群类别
        table.cell(2, 13).text = table.cell(2, 13).text.format(cause_action=self.cause_action)  # 案由
        table.cell(4, 2).text = table.cell(4, 2).text.format(apply_domicile_place=self.apply_domicile_place) #户籍所在地
        table.cell(5, 2).text = table.cell(5, 2).text.format(apply_address=self.apply_address) #现住所
        table.cell(6, 6).text = table.cell(6, 6).text.format(apply_postal_code=self.apply_postal_code) #邮政编码
        table.cell(6, 18).text = table.cell(6, 18).text.format(apply_mobile=self.apply_mobile)  # 联系电话
        table.cell(7, 2).text = table.cell(7, 2).text.format(apply_workspace=self.apply_workspace) #工作单位
        table.cell(8, 2).text = table.cell(8, 2).text.format(agent_name=self.agent_name) #代理人姓名
        table.cell(8, 11).text = table.cell(8, 11).text.format(agent_phone=self.agent_phone)  # 代理人电话
        table.cell(11, 0).text = table.cell(11, 0).text.format(description=self.description,request_items=self.request_items) #描述and请求事项和请求数额
        if len(self.apply_card) == 18:
            table.cell(3, 2).text = self.apply_card[0] #1
            table.cell(3, 3).text = self.apply_card[1]#2
            table.cell(3, 4).text = self.apply_card[2]#3
            table.cell(3, 5).text = self.apply_card[3]#4
            table.cell(3, 6).text = self.apply_card[4]#5

            table.cell(3, 8).text = self.apply_card[5]#6
            table.cell(3, 9).text = self.apply_card[6]#7
            table.cell(3, 10).text = self.apply_card[7]#8
            table.cell(3, 12).text = self.apply_card[8]#9
            table.cell(3, 15).text = self.apply_card[9]#10
            table.cell(3, 16).text = self.apply_card[10]#11

            table.cell(3, 17).text = self.apply_card[11]#12
            table.cell(3, 18).text = self.apply_card[12]#13
            table.cell(3, 19).text = self.apply_card[13]#14
            table.cell(3, 20).text = self.apply_card[14]#15
            table.cell(3, 21).text = self.apply_card[15]#16
            table.cell(3, 22).text = self.apply_card[16]#17
            table.cell(3, 23).text = self.apply_card[17]#18
        else:
            raise Exception("身份证号码格式错误")
        cell = table.cell(8,22)
        format_text = cell.text
        if self.agent_type =='法定代理人':
            #for cate in self.agent_type:
            pos = format_text.find("□" + "法定")
            format_text = format_text[:pos] + "√" + format_text[pos + 1:]
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
        elif self.agent_type == '委托代理人':
            pos = format_text.find("□" + "委托")
            format_text = format_text[:pos] + "√" + format_text[pos + 1:]
            cell.text = format_text
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            run = paragraph.add_run(format_text)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)

        temp_docx.save(filepath)


##################
# Created by ZXC #
##################
class NormalInjuryCompensationHandler(handlers.NoAuthHandler):
    """
@name 人身损害赔偿
@description 人身损害赔偿

@path
  /api/v0.1/calculator/compensation/normalinjury:
    post:
      tags:
      - "人身损害赔偿"
      description: "普通伤害赔偿"
      summary: "普通伤害赔偿"
      parameters:
      - name: "body"
        in: "body"
        description: "json形式的参数"
        required: true
        schema:
          $ref: "#/definitions/NormalInjury"
      responses:
        200:
          description: "返回成功"
          schema:
            $ref: "#/definitions/NormalInjuryResponse"
        400:
          description: "参数错误"
@endpath

@definitions
  NormalInjury:
    type: "object"
    required:
    - "region"
    - "hospital_stay"
    - "delay"
    - "victim_have_fixed_income"
    - "victim_income"
    - "victim_can_prove"
    - "victim_threeyear_annual_income"
    - "victim_job_annual_income"
    - "care_days"
    - "carer_have_fixed_income"
    - "carer_income"
    - "carer_can_prove"
    - "carer_threeyear_annual_income"
    - "local_carer_income"
    properties:
      region:
        type: "string"
        description: "受诉法院所在地"
      hospital_stay:
        type: "integer"
        description: "住院天数"
      delay:
        type: "integer"
        description: "误工期限"
      victim_have_fixed_income:
        type: "integer"
        description: "受害人有无固定收入，1为有，0为无，且和参数victim_can_prove不能同时为1"
        enum:
        - 0
        - 1
      victim_income:
        type: "integer"
        description: "受害人固定收入"
      victim_can_prove:
        type: "integer"
        description: "受害人能否举证，1为能，0为不能，且和参数victim_have_fixed_income不能同时为1"
      victim_threeyear_annual_income:
        type: "integer"
        description: "受害人近三年年均收入"
      victim_job_annual_income:
        type: "integer"
        description: "受害人所在行业年均收入"
      care_days:
        type: "integer"
        description: "护理期限"
      carer_have_fixed_income:
        type: "integer"
        description: "护理人员有无固定收入，1为有，0为无，且和参数carer_can_prove不能同时为1"
      carer_income:
        type: "integer"
        description: "护理人员固定收入"
      carer_can_prove:
        type: "integer"
        description: "护理人员能否举证，1为能，0为不能，且和参数carer_have_fixed_income不能同时为1"
      carer_threeyear_annual_income:
        type: "integer"
        description: "护理人员近三年年均收入"
      local_carer_income:
        type: "integer"
        description: "当地护工劳务报酬"
      medical_fee:
        type: "integer"
        description: "医疗费"
      traffic_fee:
        type: "integer"
        description: "交通费"
      hotel_fee:
        type: "integer"
        description: "住宿费"
      nutrition_fee:
        type: "integer"
        description: "营养费"
      identify_fee:
        type: "integer"
        description: "鉴定费"
      emotional_fee:
        type: "integer"
        description: "精神抚慰金"
      property_loss:
        type: "integer"
        description: "财产损失"

  NormalInjuryFee:
    type: "object"
    required:
    - "误工费"
    - "护理费"
    - "住院伙食补助费"
    - "财产损失"
    - "精神抚慰金"
    - "住宿费"
    - "鉴定费"
    - "医疗费"
    - "营养费"
    - "交通费"
    - "总计"
    properties:
      误工费:
        type: "integer"
      护理费:
        type: "integer"
      住院伙食补助费:
        type: "integer"
      财产损失:
        type: "integer"
      精神抚慰金:
        type: "integer"
      住宿费:
        type: "integer"
      鉴定费:
        type: "integer"
      医疗费:
        type: "integer"
      营养费:
        type: "integer"
      交通费:
        type: "integer"
      总计:
        type: "integer"

  NormalInjuryResponse:
    type: "object"
    required:
    - "code"
    - "msg"
    - "data"
    properties:
      code:
        type: "integer"
        description: "200返回成功，400参数错误，500服务器错误"
        enum:
        - 200
        - 400
        - 500
      msg:
        type: "string"
        description: "返回信息"
      data:
        $ref: "#/definitions/NormalInjuryFee"
@enddefinitions
    """
    calculator = NormalInjuryCompensation()

    @decorator.handler_except
    def prepare(self):
        super(NormalInjuryCompensationHandler, self).prepare()
        self.request_body = json.loads(self.request.body.decode())
        victim_have_fixed_income = self.request_body['victim_have_fixed_income']
        victim_can_prove = self.request_body['victim_can_prove']
        carer_have_fixed_income = self.request_body['carer_have_fixed_income']
        carer_can_prove = self.request_body['carer_can_prove']
        if victim_have_fixed_income and victim_can_prove:
            raise ArgumentTypeError('参数<victim_have_fixed_income> '
                                    '<victim_can_prove>不能同时为1')
        if carer_have_fixed_income and carer_can_prove:
            raise ArgumentTypeError('参数<carer_have_fixed_income> '
                                    '<carer_can_prove>不能同时为1')

    @decorator.threadpoll_executor
    def post(self):
        result = self.init_response_data()
        request_params = self.request_body
        params = Params(request_params)
        normal_compensation = self.calculator.calculate(params)
        result['data'] = normal_compensation
        unrequired_options = {
            '医疗费': request_params.get('medical_fee', 0),
            '交通费': request_params.get('traffic_fee', 0),
            '住宿费': request_params.get('hotel_fee', 0),
            '营养费': request_params.get('nutrition_fee', 0),
            '鉴定费': request_params.get('identify_fee', 0),
            '精神抚慰金': request_params.get('emotional_fee', 0),
            '财产损失': request_params.get('property_loss', 0),
        }
        common_fee = {}
        for key in unrequired_options.keys():
            if unrequired_options[key]:
                common_fee.update({key: unrequired_options[key]})
        result['data'].update(common_fee)
        total = sum([result['data'][key] for key in result['data'].keys()])
        result['data'].update({'总计': total})
        return result


class DisabilityCompensationHandler(handlers.NoAuthHandler):
    """
@name 人身损害赔偿
@description 人身损害赔偿

@path
  /api/v0.1/calculator/compensation/disability:
    post:
      tags:
      - "人身损害赔偿"
      description: "伤残赔偿"
      summary: "伤残赔偿"
      parameters:
      - name: "body"
        in: "body"
        description: "json形式的参数"
        required: true
        schema:
          $ref: "#/definitions/Disability"
      responses:
        200:
          description: "返回成功"
          schema:
            $ref: "#/definitions/DisabiliryResponse"
        400:
          description: "参数错误"
@endpath

@definitions
  Supported:
    type: "object"
    properties:
      supported_age:
        type: "integer"
        description: "被扶养人年龄"
      supported_is_urban:
        type: "integer"
        description: "被扶养人户口，1为城镇，0为农村"
        enum:
        - 0
        - 1
      supported_no_income:
        type: "integer"
        description: "被扶养人是否无劳动能力或无其他收入来源，1为是，0为否"
        enum:
        - 0
        - 1
      num_of_supporter:
        type: "integer"
        description: "扶养人数"

  Disability:
    type: "object"
    required:
    - "region"
    - "level"
    - "victim_is_urban"
    - "victim_age"
    - "supporteds"
    - "hospital_stay"
    - "delay"
    - "victim_have_fixed_income"
    - "victim_income"
    - "victim_can_prove"
    - "victim_threeyear_annual_income"
    - "victim_job_annual_income"
    - "care_days"
    - "carer_have_fixed_income"
    - "carer_income"
    - "carer_can_prove"
    - "carer_threeyear_annual_income"
    - "local_carer_income"
    properties:
      region:
        type: "string"
        description: "受诉法院所在地"
      level:
        type: "integer"
        description: "伤残等级"
      victim_is_urban:
        type: "integer"
        description: "是否城镇户口，1为城镇，0为农村"
        enum:
        - 0
        - 1
      victim_age:
        type: "integer"
        description: "受害人年龄"
      supporteds:
        type: "array"
        description: "被扶养人"
        items:
          $ref: "#/definitions/Supported"
      hospital_stay:
        type: "integer"
        description: "住院天数"
      delay:
        type: "integer"
        description: "误工期限"
      victim_have_fixed_income:
        type: "integer"
        description: "受害人有无固定收入，1为有，0为无，且和参数victim_can_prove不能同时为1"
        enum:
        - 0
        - 1
      victim_income:
        type: "integer"
        description: "受害人固定收入"
      victim_can_prove:
        type: "integer"
        description: "受害人能否举证，1为能，0为不能，且和参数victim_have_fixed_income不能同时为1"
      victim_threeyear_annual_income:
        type: "integer"
        description: "受害人近三年年均收入"
      victim_job_annual_income:
        type: "integer"
        description: "受害人所在行业年均收入"
      care_days:
        type: "integer"
        description: "护理期限"
      carer_have_fixed_income:
        type: "integer"
        description: "护理人员有无固定收入，1为有，0为无，且和参数carer_can_prove不能同时为1"
      carer_income:
        type: "integer"
        description: "护理人员固定收入"
      carer_can_prove:
        type: "integer"
        description: "护理人员能否举证，1为能，0为不能，且和参数carer_have_fixed_income不能同时为1"
      carer_threeyear_annual_income:
        type: "integer"
        description: "护理人员近三年年均收入"
      local_carer_income:
        type: "integer"
        description: "当地护工劳务报酬"
      medical_fee:
        type: "integer"
        description: "医疗费"
      traffic_fee:
        type: "integer"
        description: "交通费"
      hotel_fee:
        type: "integer"
        description: "住宿费"
      nutrition_fee:
        type: "integer"
        description: "营养费"
      identify_fee:
        type: "integer"
        description: "鉴定费"
      emotional_fee:
        type: "integer"
        description: "精神抚慰金"
      property_loss:
        type: "integer"
        description: "财产损失"
      equipment_fee:
        type: "integer"
        description: "残疾辅助器具费"

  DisabilityFee:
    type: "object"
    required:
    - "残疾赔偿金"
    - "被扶养人生活费"
    - "误工费"
    - "护理费"
    - "住院伙食补助费"
    - "财产损失"
    - "精神抚慰金"
    - "住宿费"
    - "鉴定费"
    - "医疗费"
    - "营养费"
    - "交通费"
    - "残疾辅助器具费"
    - "总计"
    properties:
      残疾赔偿金:
        type: "integer"
      被扶养人生活费:
        type: "integer"
      误工费:
        type: "integer"
      护理费:
        type: "integer"
      住院伙食补助费:
        type: "integer"
      财产损失:
        type: "integer"
      精神抚慰金:
        type: "integer"
      住宿费:
        type: "integer"
      鉴定费:
        type: "integer"
      医疗费:
        type: "integer"
      营养费:
        type: "integer"
      交通费:
        type: "integer"
      残疾辅助器具费:
        type: "integer"
      总计:
        type: "integer"

  DisabiliryResponse:
    type: "object"
    required:
    - "code"
    - "msg"
    - "data"
    properties:
      code:
        type: "integer"
        description: "200返回成功，400参数错误，500服务器错误"
        enum:
        - 200
        - 400
        - 500
      msg:
        type: "string"
        description: "返回信息"
      data:
        $ref: "#/definitions/DisabilityFee"
@enddefinitions
    """
    calculator = DisabilityCompensation()

    @decorator.handler_except
    def prepare(self):
        super(DisabilityCompensationHandler, self).prepare()
        self.request_body = json.loads(self.request.body.decode())
        victim_have_fixed_income = self.request_body['victim_have_fixed_income']
        victim_can_prove = self.request_body['victim_can_prove']
        carer_have_fixed_income = self.request_body['carer_have_fixed_income']
        carer_can_prove = self.request_body['carer_can_prove']
        if victim_have_fixed_income and victim_can_prove:
            raise ArgumentTypeError('参数<victim_have_fixed_income> '
                                    '<victim_can_prove>不能同时为1')
        if carer_have_fixed_income and carer_can_prove:
            raise ArgumentTypeError('参数<carer_have_fixed_income> '
                                    '<carer_can_prove>不能同时为1')

    @decorator.threadpoll_executor
    def post(self):
        result = self.init_response_data()
        request_params = self.request_body
        params = Params(request_params)
        disability_compensation = self.calculator.calculate(params)
        result['data'] = disability_compensation
        unrequired_options = {
            '医疗费': request_params.get('medical_fee', 0),
            '交通费': request_params.get('traffic_fee', 0),
            '住宿费': request_params.get('hotel_fee', 0),
            '营养费': request_params.get('nutrition_fee', 0),
            '鉴定费': request_params.get('identify_fee', 0),
            '精神抚慰金': request_params.get('emotional_fee', 0),
            '财产损失': request_params.get('property_loss', 0),
            '残疾辅助器具费': request_params.get('equipment_fee', 0)
        }
        common_fee = {}
        for key in unrequired_options.keys():
            if unrequired_options[key]:
                common_fee.update({key: unrequired_options[key]})
        result['data'].update(common_fee)
        total = sum([result['data'][key] for key in result['data'].keys()])
        result['data'].update({'总计': total})
        return result


class DeathCompensationHandler(handlers.NoAuthHandler):
    """
@name 人身损害赔偿
@description 人身损害赔偿

@path
  /api/v0.1/calculator/compensation/death:
    post:
      tags:
      - "人身损害赔偿"
      description: "死亡赔偿"
      summary: "死亡赔偿"
      parameters:
      - name: "body"
        in: "body"
        description: "json形式的参数"
        required: true
        schema:
          $ref: "#/definitions/Death"
      responses:
        200:
          description: "返回成功"
          schema:
            $ref: "#/definitions/DeathResponse"
        400:
          description: "参数错误"
@endpath

@definitions
  Death:
    type: "object"
    required:
    - "region"
    - "victim_is_urban"
    - "victim_age"
    - "supporteds"
    - "hospital_stay"
    - "delay"
    - "victim_have_fixed_income"
    - "victim_income"
    - "victim_can_prove"
    - "victim_threeyear_annual_income"
    - "victim_job_annual_income"
    - "care_days"
    - "carer_have_fixed_income"
    - "carer_income"
    - "carer_can_prove"
    - "carer_threeyear_annual_income"
    - "local_carer_income"
    properties:
      region:
        type: "string"
        description: "受诉法院所在地"
      victim_is_urban:
        type: "integer"
        description: "是否城镇户口，1为城镇，0为农村"
        enum:
        - 0
        - 1
      victim_age:
        type: "integer"
        description: "受害人年龄"
      supporteds:
        type: "array"
        description: "被扶养人"
        items:
          $ref: "#/definitions/Supported"
      hospital_stay:
        type: "integer"
        description: "住院天数"
      delay:
        type: "integer"
        description: "误工期限"
      victim_have_fixed_income:
        type: "integer"
        description: "受害人有无固定收入，1为有，0为无，且和参数victim_can_prove不能同时为1"
        enum:
        - 0
        - 1
      victim_income:
        type: "integer"
        description: "受害人固定收入"
      victim_can_prove:
        type: "integer"
        description: "受害人能否举证，1为能，0为不能，且和参数victim_have_fixed_income不能同时为1"
      victim_threeyear_annual_income:
        type: "integer"
        description: "受害人近三年年均收入"
      victim_job_annual_income:
        type: "integer"
        description: "受害人所在行业年均收入"
      care_days:
        type: "integer"
        description: "护理期限"
      carer_have_fixed_income:
        type: "integer"
        description: "护理人员有无固定收入，1为有，0为无，且和参数carer_can_prove不能同时为1"
      carer_income:
        type: "integer"
        description: "护理人员固定收入"
      carer_can_prove:
        type: "integer"
        description: "护理人员能否举证，1为能，0为不能，且和参数carer_have_fixed_income不能同时为1"
      carer_threeyear_annual_income:
        type: "integer"
        description: "护理人员近三年年均收入"
      local_carer_income:
        type: "integer"
        description: "当地护工劳务报酬"
      medical_fee:
        type: "integer"
        description: "医疗费"
      traffic_fee:
        type: "integer"
        description: "交通费"
      hotel_fee:
        type: "integer"
        description: "住宿费"
      nutrition_fee:
        type: "integer"
        description: "营养费"
      identify_fee:
        type: "integer"
        description: "鉴定费"
      emotional_fee:
        type: "integer"
        description: "精神抚慰金"
      property_loss:
        type: "integer"
        description: "财产损失"

  DeathFee:
    type: "object"
    required:
    - "死亡补偿费"
    - "被扶养人生活费"
    - "误工费"
    - "护理费"
    - "住院伙食补助费"
    - "财产损失"
    - "精神抚慰金"
    - "住宿费"
    - "鉴定费"
    - "医疗费"
    - "营养费"
    - "交通费"
    - "丧葬费"
    - "总计"
    properties:
      死亡补偿费:
        type: "integer"
      被扶养人生活费:
        type: "integer"
      误工费:
        type: "integer"
      护理费:
        type: "integer"
      住院伙食补助费:
        type: "integer"
      财产损失:
        type: "integer"
      精神抚慰金:
        type: "integer"
      住宿费:
        type: "integer"
      鉴定费:
        type: "integer"
      医疗费:
        type: "integer"
      营养费:
        type: "integer"
      交通费:
        type: "integer"
      丧葬费:
        type: "integer"
      总计:
        type: "integer"

  DeathResponse:
    type: "object"
    required:
    - "code"
    - "msg"
    - "data"
    properties:
      code:
        type: "integer"
        description: "200返回成功，400参数错误，500服务器错误"
        enum:
        - 200
        - 400
        - 500
      msg:
        type: "string"
        description: "返回信息"
      data:
        $ref: "#/definitions/DeathFee"
@enddefinitions
    """
    calculator = DeathCompensation()

    @decorator.handler_except
    def prepare(self):
        super(DeathCompensationHandler, self).prepare()
        self.request_body = json.loads(self.request.body.decode())
        victim_have_fixed_income = self.request_body['victim_have_fixed_income']
        victim_can_prove = self.request_body['victim_can_prove']
        carer_have_fixed_income = self.request_body['carer_have_fixed_income']
        carer_can_prove = self.request_body['carer_can_prove']
        if victim_have_fixed_income and victim_can_prove:
            raise ArgumentTypeError('参数<victim_have_fixed_income> '
                                    '<victim_can_prove>不能同时为1')
        if carer_have_fixed_income and carer_can_prove:
            raise ArgumentTypeError('参数<carer_have_fixed_income> '
                                    '<carer_can_prove>不能同时为1')

    @decorator.threadpoll_executor
    def post(self):
        result = self.init_response_data()
        request_params = self.request_body
        request_params.update({'level': 1})
        params = Params(request_params)
        death_compensation = self.calculator.calculate(params)
        result['data'] = death_compensation
        unrequired_options = {
            '医疗费': request_params.get('medical_fee', 0),
            '交通费': request_params.get('traffic_fee', 0),
            '住宿费': request_params.get('hotel_fee', 0),
            '营养费': request_params.get('nutrition_fee', 0),
            '鉴定费': request_params.get('identify_fee', 0),
            '精神抚慰金': request_params.get('emotional_fee', 0),
            '财产损失': request_params.get('property_loss', 0),
        }
        common_fee = {}
        for key in unrequired_options.keys():
            if unrequired_options[key]:
                common_fee.update({key: unrequired_options[key]})
        result['data'].update(common_fee)
        total = sum([result['data'][key] for key in result['data'].keys()])
        result['data'].update({'总计': total})
        return result


handlers = [
    (r"/intelligentpretrial/android/law_push/Unsigned_Contract_Handler", CalculatorHandler),
    (r"/intelligentpretrial/android/law_push/lawapply", LawapplyHandler),
    (r"/api/android/law_push/ZNYS_Labor_Handler", ZNYS_Labor_Handler),
    (r"/intelligentpretrial/android/law_push/aidform", LawAidFormHandler),

    (r'/api/v0.1/calculator/compensation/normalinjury', NormalInjuryCompensationHandler),
    (r'/api/v0.1/calculator/compensation/disability', DisabilityCompensationHandler),
    (r'/api/v0.1/calculator/compensation/death', DeathCompensationHandler)
]
